import os
import requests
import pandas as pd
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, filters, CallbackQueryHandler, CallbackContext
from docx import Document

# Set up API keys from environment variables
TELEGRAM_GOOGLE_BOOKS_API_KEY = '7124925550:AAEx8fQtFQBI9HQ9Qe1qUE05Y_2zlG_FxGA'
GOOGLE_BOOKS_API_KEY = 'AIzaSyD2HipitHKX4zc5Sc1xx4CMoBCMpjqCUBk'

# Initialize reading list document
READING_LIST_DOC = 'reading_list.docx'
if not os.path.exists(READING_LIST_DOC):
    doc = Document()
    doc.add_heading('My Reading List', 0)
    doc.save(READING_LIST_DOC)

async def start(update: Update, context: CallbackContext) -> None:
    await update.message.reply_text("Welcome to PagePal! Use /help to see available commands.")

async def help_command(update: Update, context: CallbackContext) -> None:
    commands = (
        "/start - Welcome message\n"
        "/book - Search for books by genre\n"
        "/preview - Get a preview link for a specific book\n"
        "/list - Start managing your reading list\n"
        "/reading_list - Manage your reading list (add, delete, view)\n"
        "/help - Show this help message"
    )
    await update.message.reply_text(commands)

async def book(update: Update, context: CallbackContext) -> None:
    await update.message.reply_text("Please enter a genre (e.g., Fiction, Science, History):")

def fetch_books_by_genre(genre: str, GOOGLE_BOOKS_API_KEY: str):
    try:
        url = f"https://www.googleapis.com/books/v1/volumes?q=subject:{genre}&key={GOOGLE_BOOKS_API_KEY}&maxResults=10"
        response = requests.get(url)
        response.raise_for_status()
        books = response.json().get('items', [])
        
        if not books:
            return None

        book_data = []
        for book in books:
            info = book['volumeInfo']
            book_data.append({
                'Title': info.get('title'),
                'Author': ', '.join(info.get('authors', [])),
                'Description': info.get('description', 'No description available'),
                'Published Year': info.get('publishedDate', 'N/A').split('-')[0],
                'Language': info.get('language', 'N/A'),
                'Preview Link': info.get('previewLink', 'N/A')
            })

        df = pd.DataFrame(book_data)
        filename = f'{genre}_books.csv'
        df.to_csv(filename, index=False)
        return filename
    except requests.RequestException as e:
        print(f"Request failed: {e}")
        return None

async def send_book_list(update: Update, context: CallbackContext) -> None:
    genre = update.message.text
    csv_file = fetch_books_by_genre(genre,GOOGLE_BOOKS_API_KEY)
    
    if csv_file:
        await update.message.reply_text(f"Here are some books in the {genre} genre:")
        await update.message.reply_document(open(csv_file, 'rb'))
    else:
        await update.message.reply_text(f"Sorry, no books found in the {genre} genre.")

async def preview(update: Update, context: CallbackContext) -> None:
    await update.message.reply_text("Please enter the book name you want a preview for:")

def fetch_book_preview(book_name: str, GOOGLE_BOOKS_API_KEY: str):
    try:
        url = f"https://www.googleapis.com/books/v1/volumes?q=intitle:{book_name}&key={GOOGLE_BOOKS_API_KEY}&maxResults=1"
        response = requests.get(url)
        response.raise_for_status()
        books = response.json().get('items', [])
        
        if books:
            preview_link = books[0]['volumeInfo'].get('previewLink', 'No preview available')
            return preview_link
        return "No preview available"
    except requests.RequestException as e:
        print(f"Request failed: {e}")
        return "No preview available"

async def send_book_preview(update: Update, context: CallbackContext) -> None:
    book_name = update.message.text
    preview_link = fetch_book_preview(book_name,GOOGLE_BOOKS_API_KEY)
    
    await update.message.reply_text(f"Preview for '{book_name}': {preview_link}")

async def list_books(update: Update, context: CallbackContext) -> None:
    await update.message.reply_text("Enter the book name and use /reading_list to manage your reading list.")

async def reading_list(update: Update, context: CallbackContext) -> None:
    keyboard = [
        [InlineKeyboardButton("Add a Book", callback_data='add')],
        [InlineKeyboardButton("Delete a Book", callback_data='delete')],
        [InlineKeyboardButton("View Reading List", callback_data='view')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text("Manage your reading list:", reply_markup=reply_markup)

def add_book_to_reading_list(title: str, preview_link: str):
    doc = Document(READING_LIST_DOC)
    doc.add_paragraph(f"Title: {title}\nPreview: {preview_link}\n")
    doc.save(READING_LIST_DOC)

def delete_book_from_reading_list(title: str):
    doc = Document(READING_LIST_DOC)
    paragraphs = doc.paragraphs
    for i, para in enumerate(paragraphs):
        if title in para.text:
            paragraphs[i].clear()  # Clear the paragraph that contains the title
            break
    doc.save(READING_LIST_DOC)

async def button(update: Update, context: CallbackContext) -> None: nbbb
    query = update.callback_query
    await query.answer()

    if query.data == 'add':
        await query.edit_message_text(text="Please enter the book title and preview link (comma-separated) to add:")
        context.user_data['action'] = 'add'

    elif query.data == 'delete':
        await query.edit_message_text(text="Please enter the book title to delete:")
        context.user_data['action'] = 'delete'

    elif query.data == 'view':
        await query.edit_message_text(text="Here is your reading list:")
        await query.message.reply_document(open(READING_LIST_DOC, 'rb'))

async def handle_text(update: Update, context: CallbackContext) -> None:
    action = context.user_data.get('action')

    if action == 'add':
        try:
            title, preview_link = update.message.text.split(',', 1)
            add_book_to_reading_list(title.strip(), preview_link.strip())
            await update.message.reply_text(f"Book '{title}' added to the reading list.")
        except ValueError:
            await update.message.reply_text("Invalid format. Please enter the title and preview link separated by a comma.")
        context.user_data['action'] = None

    elif action == 'delete':
        title = update.message.text.strip()
        delete_book_from_reading_list(title)
        await update.message.reply_text(f"Book '{title}' deleted from the reading list.")
        context.user_data['action'] = None

def main():
    application = Application.builder().token(TELEGRAM_GOOGLE_BOOKS_API_KEY).build()

    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("help", help_command))
    application.add_handler(CommandHandler("book", book))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, send_book_list, block=False))
    application.add_handler(CommandHandler("preview", preview))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, send_book_preview, block=False))
    application.add_handler(CommandHandler("list", list_books))
    application.add_handler(CommandHandler("reading_list", reading_list))
    application.add_handler(CallbackQueryHandler(button))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))

    application.run_polling()

if __name__ == '__main__':
    main()
